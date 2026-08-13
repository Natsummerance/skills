#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown 转 DOCX 工具
支持：标题、表格、粗体/斜体/行内代码、有序/无序列表、代码块、水平线

依赖：pip install python-docx
"""

import sys
import os
import re
import argparse
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def add_formatted_text(paragraph, text: str, base_font_name: str = 'Microsoft YaHei',
                       base_font_size: int = 11):
    """
    解析行内 Markdown 格式并添加到段落中。
    支持：**粗体**、*斜体*、`行内代码`
    """
    # 用正则拆分文本，保留格式标记
    # 匹配顺序：**粗体** > *斜体* > `代码`
    pattern = r'(\*\*.*?\*\*|`.*?`|\*.*?\*)'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue

        if part.startswith('**') and part.endswith('**'):
            # 粗体
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
            run.font.name = base_font_name
            run.font.size = Pt(base_font_size)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), base_font_name)

        elif part.startswith('`') and part.endswith('`'):
            # 行内代码
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(base_font_size - 1)
            # 灰色背景通过高亮实现
            run.font.highlight_color = 15  # wdColorGray15

        elif part.startswith('*') and part.endswith('*'):
            # 斜体
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
            run.font.name = base_font_name
            run.font.size = Pt(base_font_size)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), base_font_name)

        else:
            # 普通文本
            run = paragraph.add_run(part)
            run.font.name = base_font_name
            run.font.size = Pt(base_font_size)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), base_font_name)


def parse_heading_level(line: str) -> int:
    """解析标题级别，返回 0 表示不是标题"""
    match = re.match(r'^(#{1,6})\s+(.*)', line)
    if match:
        return len(match.group(1))
    return 0


def parse_markdown(md_content: str) -> list:
    """解析 Markdown 内容为结构化数据"""
    lines = md_content.split('\n')
    elements = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # 标题
        level = parse_heading_level(line)
        if level > 0:
            text = re.sub(r'^#{1,6}\s+', '', line).strip()
            elements.append({'type': 'heading', 'level': level, 'text': text})
            i += 1
            continue

        # 代码块
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # 跳过结束的 ```
            elements.append({'type': 'code', 'content': '\n'.join(code_lines)})
            continue

        # 表格
        if '|' in line and i + 1 < len(lines) and re.match(r'^\s*\|?\s*[-:]+', lines[i + 1]):
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            elements.append({'type': 'table', 'lines': table_lines})
            continue

        # 无序列表
        if re.match(r'^\s*[-*]\s+', line):
            list_items = []
            while i < len(lines) and re.match(r'^\s*[-*]\s+', lines[i]):
                text = re.sub(r'^\s*[-*]\s+', '', lines[i]).strip()
                list_items.append(text)
                i += 1
            elements.append({'type': 'list', 'ordered': False, 'items': list_items})
            continue

        # 有序列表
        if re.match(r'^\s*\d+\.\s+', line):
            list_items = []
            while i < len(lines) and re.match(r'^\s*\d+\.\s+', lines[i]):
                text = re.sub(r'^\s*\d+\.\s+', '', lines[i]).strip()
                list_items.append(text)
                i += 1
            elements.append({'type': 'list', 'ordered': True, 'items': list_items})
            continue

        # 水平线
        if re.match(r'^\s*(---|\*\*\*|___)\s*$', line):
            elements.append({'type': 'hr'})
            i += 1
            continue

        # 段落
        if line.strip():
            elements.append({'type': 'paragraph', 'text': line.strip()})

        i += 1

    return elements


def add_table_to_doc(doc: Document, table_lines: list):
    """添加表格到文档，带样式"""
    rows = []
    for line in table_lines:
        # 跳过分隔行
        if re.match(r'^\s*\|?\s*[-:]+', line):
            continue
        cells = [cell.strip() for cell in line.split('|')]
        # 去掉首尾空元素（由首尾 | 产生）
        if cells and cells[0] == '':
            cells = cells[1:]
        if cells and cells[-1] == '':
            cells = cells[:-1]
        if cells:
            rows.append(cells)

    if not rows:
        return

    num_cols = max(len(row) for row in rows)
    # 补齐列数
    for row in rows:
        while len(row) < num_cols:
            row.append('')

    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            # 清空默认段落
            cell.paragraphs[0].clear()
            add_formatted_text(cell.paragraphs[0], cell_text, base_font_size=10)

            if i == 0:
                # 表头：加粗 + 蓝色背景 + 白色文字
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                # 设置背景色
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="3498DB"/>')
                cell._element.get_or_add_tcPr().append(shading)
            else:
                # 交替行背景色
                if i % 2 == 0:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0"/>')
                    cell._element.get_or_add_tcPr().append(shading)


def add_horizontal_rule(doc: Document):
    """添加水平线（段落底部边框）"""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def add_code_block(doc: Document, code: str):
    """添加代码块（等宽字体 + 灰色背景）"""
    for line in code.split('\n'):
        p = doc.add_paragraph()
        # 段落背景色
        pPr = p._element.get_or_add_pPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F4F4"/>')
        pPr.append(shading)
        # 左缩进
        ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="360"/>')
        pPr.append(ind)

        run = p.add_run(line if line else ' ')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def elements_to_docx(elements: list, output_path: str):
    """将结构化元素转换为 DOCX"""
    doc = Document()

    # 设置默认字体（中文回退）
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # 配置标题样式的中文字体
    for i in range(1, 5):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Microsoft YaHei'
        heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    for elem in elements:
        elem_type = elem['type']

        if elem_type == 'heading':
            level = min(elem['level'], 4)
            text = elem['text']
            heading = doc.add_heading(level=level)
            add_formatted_text(heading, text, base_font_size=11 + (5 - level) * 2)

        elif elem_type == 'paragraph':
            p = doc.add_paragraph()
            add_formatted_text(p, elem['text'])

        elif elem_type == 'list':
            for item in elem['items']:
                if elem['ordered']:
                    p = doc.add_paragraph(style='List Number')
                else:
                    p = doc.add_paragraph(style='List Bullet')
                p.clear()
                add_formatted_text(p, item)

        elif elem_type == 'table':
            add_table_to_doc(doc, elem['lines'])
            doc.add_paragraph()  # 表后空行

        elif elem_type == 'code':
            add_code_block(doc, elem['content'])

        elif elem_type == 'hr':
            add_horizontal_rule(doc)

    doc.save(output_path)


def md_to_docx(input_path: str, output_path: str = None) -> str:
    """将 Markdown 文件转换为 DOCX"""
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"输入文件不存在: {input_path}")

    if output_path is None:
        base = os.path.splitext(input_path)[0]
        output_path = f"{base}.docx"

    with open(input_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    elements = parse_markdown(md_content)
    elements_to_docx(elements, output_path)

    return output_path


def main():
    parser = argparse.ArgumentParser(description="Markdown 转 DOCX 工具（支持中文、表格、粗体）")
    parser.add_argument("input", help="输入的 Markdown 文件路径")
    parser.add_argument("-o", "--output", help="输出的 DOCX 文件路径（默认与输入同名）")
    args = parser.parse_args()

    try:
        result = md_to_docx(args.input, args.output)
        print(f"[SUCCESS] DOCX 生成成功: {result}")
    except Exception as e:
        print(f"[FAILED] 转换失败: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
